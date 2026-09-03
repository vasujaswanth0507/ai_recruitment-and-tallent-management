import io
import re
import logging
from typing import Union, BinaryIO

# Set up logging
logger = logging.getLogger(__name__)


def clean_extracted_text(text: str) -> str:
    """Normalizes whitespace and cleans up raw extracted text."""
    if not text:
        return ""
    # Normalize carriage returns and newlines
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Replace multiple consecutive spaces/tabs with single space
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    # Join with newlines, filtering out excessive blank lines (max 2 consecutive newlines)
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_text_from_pdf(file_input: Union[bytes, BinaryIO, str]) -> str:
    """
    Extracts text from PDF file bytes or path.
    Primary method: pdfplumber.
    Fallback method: pypdf / PyPDF2.
    Handles encrypted, image-only, or corrupted files gracefully.
    """
    pdf_bytes = None
    if isinstance(file_input, bytes):
        pdf_bytes = file_input
    elif isinstance(file_input, str):
        with open(file_input, "rb") as f:
            pdf_bytes = f.read()
    elif hasattr(file_input, "read"):
        pdf_bytes = file_input.read()
        if hasattr(file_input, "seek"):
            file_input.seek(0)
    else:
        raise ValueError("Invalid input for PDF extraction. Expected bytes, path, or file stream.")

    extracted_text = ""

    # Primary Attempt: pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            pages_text = []
            for i, page in enumerate(pdf.pages):
                page_txt = page.extract_text()
                if page_txt:
                    pages_text.append(page_txt)

            extracted_text = "\n".join(pages_text)
    except Exception as e1:
        logger.warning(f"pdfplumber failed: {e1}. Attempting PyPDF fallback...")

    # Fallback Attempt: pypdf / PyPDF2
    if not extracted_text.strip():
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    raise ValueError("The PDF document is password protected / encrypted.")

            pages_text = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt)
            extracted_text = "\n".join(pages_text)
        except Exception as e2:
            logger.error(f"PyPDF extraction also failed: {e2}")
            if "password" in str(e1).lower() or "password" in str(e2).lower():
                raise ValueError("PDF is encrypted and cannot be parsed without password.")
            elif not extracted_text.strip():
                raise ValueError(
                    "Could not extract text from PDF. The file may be scanned/image-only or corrupted."
                )

    cleaned = clean_extracted_text(extracted_text)
    if not cleaned:
        raise ValueError(
            "Extracted text is empty. File might be an image-only scanned document or empty."
        )

    return cleaned


def extract_text_from_docx(file_input: Union[bytes, BinaryIO, str]) -> str:
    """
    Extracts text from DOCX file bytes or path using python-docx.
    """
    try:
        import docx

        if isinstance(file_input, bytes):
            stream = io.BytesIO(file_input)
        elif isinstance(file_input, str):
            stream = file_input
        elif hasattr(file_input, "read"):
            stream = io.BytesIO(file_input.read())
            if hasattr(file_input, "seek"):
                file_input.seek(0)
        else:
            raise ValueError("Invalid input for DOCX extraction.")

        doc = docx.Document(stream)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        extracted = "\n".join(full_text)
        cleaned = clean_extracted_text(extracted)
        if not cleaned:
            raise ValueError("DOCX document contains no readable text.")
        return cleaned
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")


def parse_document(file_name: str, file_bytes: bytes) -> str:
    """
    Universal document parser routing by file extension.
    Supported extensions: .pdf, .docx, .doc, .txt
    """
    filename_lower = file_name.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        return clean_extracted_text(file_bytes.decode("utf-8", errors="ignore"))
    else:
        raise ValueError(
            f"Unsupported file format '{file_name}'. Supported formats are: .pdf, .docx, .txt"
        )
